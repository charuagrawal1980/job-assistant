import logging
from datetime import datetime

import pdfplumber
import pandas as pd
import mammoth
import markdownify
import markdown


from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def get_wordfile_markdown(filename)->str:
    with open(filename, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value
        markdown = markdownify.markdownify(html)
        return markdown
    
def save_markdowntext_to_word(markdown_text, filename):
    
    html_content = markdown.markdown(markdown_text)
    doc = Document()
    add_html_to_doc(doc, html_content)
    output_file_path = f"{filename}.docx" 
    path = doc.save(output_file_path)
    return output_file_path

def save_resume_to_pdf(text: str) -> str:

    if not text:
        return None
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    lines = text.split('\n')
    for line in lines:
        pdf.multi_cell(0, 10, txt=line)
    
    filename = f"tailored_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    pdf.output(filename)
    return filename

def get_resume_text(resume_filename: str) -> str:

    resume_text = ""
    try:
        if resume_filename.endswith(".pdf"):
            with pdfplumber.open(resume_filename) as pdf:
                resume_text = " ".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
        else:
            with open(resume_filename, "r") as file:
                resume_text = file.read()
        return resume_text
    except Exception as e:
        logger.error(f"Error extracting resume text: {str(e)}", exc_info=True)
        return ""
    
def get_jobs_from_excel(fileame):
        df = pd.read_excel(fileame)
        json_data = []
    
        for _, row in df.iterrows():
            job_entry = {
            "company_name": str(row['company_name']).strip(),
            "job_location": "",
            "job_description": str(row['job_description']).strip(),
            "job_title":row['job_title'],
            "job_url": row['job_url']
        }
            json_data.append(job_entry)
    
        return json_data

def add_html_to_doc(doc, html_content):

    """Parses HTML content and adds it to the Word document with appropriate formatting."""
    soup = BeautifulSoup(html_content, "html.parser")
    
    def process_element(element, parent_paragraph=None):
        if isinstance(element, str):
            if parent_paragraph:
                parent_paragraph.add_run(element)
            else:
                doc.add_paragraph(element)
            return

        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
            font_size = {
                "h1": 16, "h2": 14, "h3": 12,
                "h4": 11, "h5": 10, "h6": 10
            }
            run.font.size = Pt(font_size.get(element.name, 10))
            
        elif element.name == "p":
            p = doc.add_paragraph()
            for child in element.children:
                process_element(child, p)
                
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                for child in li.children:
                    process_element(child, p)
                    
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Number")
                for child in li.children:
                    process_element(child, p)
                    
        elif element.name in ["strong", "b"]:
            run = parent_paragraph.add_run(element.get_text())
            run.bold = True
            
        elif element.name in ["em", "i"]:
            run = parent_paragraph.add_run(element.get_text())
            run.italic = True
            
        elif element.name == "a":
            run = parent_paragraph.add_run(f"{element.get_text()} ({element.get('href', '')})")
            run.underline = True
            
        else:
            # Process all child elements
            for child in element.children:
                process_element(child)

    # Process the document
    for element in soup.body.children if soup.body else soup.children:
        if isinstance(element, str) and element.strip() == "":
            continue
        process_element(element)

    return doc

def download_all_resumes():
    return ""